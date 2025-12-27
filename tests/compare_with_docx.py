"""
Script to compare CSV export with Word document
"""

import pandas as pd
import sys
import io
from pathlib import Path

# Set UTF-8 encoding for output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Paths
csv_path = Path(__file__).parent / "2025-12-27T09-34_export.csv"
docx_path = Path(__file__).parent.parent.parent / "5 TÀI LIỆU CHÍNH" / "CHECK VN10.docx"

print("=" * 80)
print("SO SANH CSV EXPORT VOI WORD DOCUMENT")
print("=" * 80)

# Read CSV
print("\n1. DOC FILE CSV:")
print(f"   Path: {csv_path}")
try:
    df_csv = pd.read_csv(csv_path)
    print(f"   OK: Doc thanh cong!")
    print(f"\n   Noi dung CSV:")
    print(df_csv.to_string(index=False))
except Exception as e:
    print(f"   LOI: {e}")
    sys.exit(1)

# Try to read DOCX
print("\n2. DOC FILE WORD:")
print(f"   Path: {docx_path}")
try:
    import docx
    doc = docx.Document(docx_path)
    
    print(f"   OK: Doc thanh cong!")
    print(f"\n   Noi dung Word Document:")
    print("-" * 80)
    
    # Extract text from all paragraphs
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            print(para.text)
    
    # Extract tables
    print("\n   Tables trong document:")
    for i, table in enumerate(doc.tables):
        print(f"\n   Table {i+1}:")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            print("   | " + " | ".join(row_data) + " |")
    
except ImportError:
    print("   CAN CAI DAT python-docx:")
    print("   pip install python-docx")
    print("\n   Thay vao do, hien thi noi dung CSV:")
    print(df_csv.to_string(index=False))
except Exception as e:
    print(f"   LOI: {e}")
    print("\n   Thay vao do, hien thi noi dung CSV:")
    print(df_csv.to_string(index=False))

print("\n" + "=" * 80)
print("HOAN TAT")
print("=" * 80)

